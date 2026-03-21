import os
import pickle
import numpy as np
import streamlit as st
from langchain_core.documents import Document
from langchain_core.embeddings import Embeddings
from langchain_community.vectorstores import FAISS
from langchain_community.document_loaders import TextLoader, PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from sklearn.feature_extraction.text import TfidfVectorizer

FAISS_INDEX_PATH = "faiss_index"

# ── TF-IDF Embeddings — no download, works everywhere ──
class TFIDFEmbeddings(Embeddings):
    def __init__(self):
        self.vectorizer = TfidfVectorizer(
            max_features=384,
            ngram_range=(1, 2),
            sublinear_tf=True
        )
        self.fitted = False
        self.dim = 384

    def _normalize(self, vec):
        norm = np.linalg.norm(vec)
        return (vec / norm).tolist() if norm > 0 else vec.tolist()

    def fit(self, texts):
        self.vectorizer.fit(texts)
        self.fitted = True

    def embed_documents(self, texts):
        if not self.fitted:
            self.fit(texts)
        matrix = self.vectorizer.transform(texts).toarray()
        return [self._normalize(row) for row in matrix]

    def embed_query(self, text):
        if not self.fitted:
            return [0.0] * self.dim
        vec = self.vectorizer.transform([text]).toarray()[0]
        return self._normalize(vec)

# Global embeddings instance (keeps vectorizer fitted across calls)
@st.cache_resource
def get_embeddings():
    return TFIDFEmbeddings()

def get_splitter():
    return RecursiveCharacterTextSplitter(
        chunk_size=500,
        chunk_overlap=100,
        separators=["\n\n", "\n", ".", "!", "?", " ", ""]
    )

def load_documents():
    docs = []

    if not os.path.exists("data"):
        os.makedirs("data")

    if os.path.exists("data/sample.txt"):
        try:
            loader = TextLoader("data/sample.txt", encoding="utf-8")
            docs.extend(loader.load())
        except Exception as e:
            print(f"Could not load sample.txt: {e}")

    for file in os.listdir("data"):
        if file.endswith(".pdf"):
            try:
                loader = PyPDFLoader(os.path.join("data", file))
                docs.extend(loader.load())
            except Exception as e:
                print(f"Could not load {file}: {e}")

    if not docs:
        docs = [Document(
            page_content="This is LinguaBot, a multilingual AI assistant. Upload your documents to get started.",
            metadata={"source": "default"}
        )]

    return docs

def create_vectorstore():
    documents  = load_documents()
    splitter   = get_splitter()
    chunks     = splitter.split_documents(documents)
    embeddings = get_embeddings()

    # Fit vectorizer on all document texts
    all_texts = [doc.page_content for doc in chunks]
    embeddings.fit(all_texts)

    db = FAISS.from_documents(chunks, embeddings)

    try:
        db.save_local(FAISS_INDEX_PATH)
        # Save fitted vectorizer alongside index
        with open(f"{FAISS_INDEX_PATH}/vectorizer.pkl", "wb") as f:
            pickle.dump(embeddings.vectorizer, f)
        print(f"Vectorstore saved to '{FAISS_INDEX_PATH}'")
    except Exception as e:
        print(f"Could not save index: {e}")

    return db

def load_vectorstore():
    embeddings = get_embeddings()

    if os.path.exists(FAISS_INDEX_PATH) and os.path.exists(f"{FAISS_INDEX_PATH}/vectorizer.pkl"):
        try:
            print("Loading vectorstore from disk...")
            # Restore fitted vectorizer
            with open(f"{FAISS_INDEX_PATH}/vectorizer.pkl", "rb") as f:
                embeddings.vectorizer = pickle.load(f)
                embeddings.fitted = True

            db = FAISS.load_local(
                FAISS_INDEX_PATH,
                embeddings,
                allow_dangerous_deserialization=True
            )
            return db
        except Exception as e:
            print(f"Could not load saved index, rebuilding: {e}")

    print("Building vectorstore from documents...")
    return create_vectorstore()

def merge_documents_into_db(db, file_bytes, filename):
    import tempfile

    ext  = filename.lower().split(".")[-1]
    docs = []

    with tempfile.NamedTemporaryFile(delete=False, suffix=f".{ext}") as tmp:
        tmp.write(file_bytes)
        tmp_path = tmp.name

    try:
        if ext == "pdf":
            loader = PyPDFLoader(tmp_path)
            docs   = loader.load()
        elif ext == "txt":
            loader = TextLoader(tmp_path, encoding="utf-8")
            docs   = loader.load()
        elif ext == "docx":
            from docx import Document as DocxDocument
            docx      = DocxDocument(tmp_path)
            full_text = "\n".join([p.text for p in docx.paragraphs if p.text.strip()])
            docs      = [Document(page_content=full_text, metadata={"source": filename})]
        else:
            raise ValueError(f"Unsupported file type: {ext}")
    finally:
        os.unlink(tmp_path)

    if not docs:
        return db, 0

    for doc in docs:
        doc.metadata["source"] = filename

    splitter   = get_splitter()
    chunks     = splitter.split_documents(docs)
    embeddings = get_embeddings()

    # Refit vectorizer with new texts added
    new_texts = [doc.page_content for doc in chunks]
    embeddings.fit(new_texts)

    new_db = FAISS.from_documents(chunks, embeddings)
    db.merge_from(new_db)

    try:
        db.save_local(FAISS_INDEX_PATH)
        with open(f"{FAISS_INDEX_PATH}/vectorizer.pkl", "wb") as f:
            pickle.dump(embeddings.vectorizer, f)
    except Exception as e:
        print(f"Could not save after merge: {e}")

    return db, len(chunks)