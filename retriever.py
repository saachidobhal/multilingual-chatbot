import os
import re
from langchain_community.document_loaders import TextLoader, PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document


def get_splitter():
    return RecursiveCharacterTextSplitter(
        chunk_size=500,
        chunk_overlap=100,
        separators=["\n\n", "\n", ".", "!", "?", " ", ""]
    )


def load_documents():
    docs = []

    if not os.path.exists("data"):
        os.makedirs("data")

    if os.path.exists("data/sample.txt"):
        try:
            loader = TextLoader("data/sample.txt", encoding="utf-8")
            docs.extend(loader.load())
        except Exception as e:
            print(f"Could not load sample.txt: {e}")

    for file in os.listdir("data"):
        if file.endswith(".pdf"):
            try:
                loader = PyPDFLoader(os.path.join("data", file))
                docs.extend(loader.load())
            except Exception as e:
                print(f"Could not load {file}: {e}")

    if not docs:
        docs = [Document(
            page_content="This is LinguaBot, a multilingual AI assistant. Upload documents to get started.",
            metadata={"source": "default"}
        )]

    return docs


class DocStore:
    """Zero-dependency document store using keyword search. No FAISS, no embeddings."""

    def __init__(self):
        self.chunks = []

    def build(self):
        documents = load_documents()
        splitter  = get_splitter()
        self.chunks = splitter.split_documents(documents)
        print(f"DocStore built with {len(self.chunks)} chunks")
        return self

    def search(self, query, k=3):
        if not self.chunks:
            return []
        query_words = set(re.findall(r'\w+', query.lower()))
        scored = []
        for chunk in self.chunks:
            text_words = set(re.findall(r'\w+', chunk.page_content.lower()))
            score = len(query_words & text_words)
            scored.append((score, chunk))
        scored.sort(key=lambda x: x[0], reverse=True)
        return [chunk for _, chunk in scored[:k]]

    def add(self, file_bytes, filename):
        import tempfile
        ext  = filename.lower().split(".")[-1]
        docs = []

        with tempfile.NamedTemporaryFile(delete=False, suffix=f".{ext}") as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name

        try:
            if ext == "pdf":
                loader = PyPDFLoader(tmp_path)
                docs   = loader.load()
            elif ext == "txt":
                loader = TextLoader(tmp_path, encoding="utf-8")
                docs   = loader.load()
            elif ext == "docx":
                from docx import Document as DocxDocument
                docx      = DocxDocument(tmp_path)
                full_text = "\n".join([p.text for p in docx.paragraphs if p.text.strip()])
                docs      = [Document(page_content=full_text, metadata={"source": filename})]
        finally:
            os.unlink(tmp_path)

        if not docs:
            return 0

        for doc in docs:
            doc.metadata["source"] = filename

        splitter = get_splitter()
        chunks   = splitter.split_documents(docs)
        self.chunks.extend(chunks)
        return len(chunks)
